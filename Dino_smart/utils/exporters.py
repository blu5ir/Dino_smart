import io
from fpdf import FPDF
from docx import Document as DocxDoc
from docx.shared import Inches, Pt, RGBColor

class StudyBuddyExporter:
    @staticmethod
    def export_to_markdown(title: str, content_dict: dict) -> str:
        """Converts study guide sections or revision notes into raw Markdown."""
        lines = [f"# {title}", ""]
        for section, content in content_dict.items():
            lines.append(f"## {section}")
            lines.append(content)
            lines.append("")
        return "\n".join(lines)

    @staticmethod
    def export_to_html(title: str, content_dict: dict) -> str:
        """Generates a premium, styled HTML document for study guides."""
        md = StudyBuddyExporter.export_to_markdown(title, content_dict)
        # Convert markdown to basic HTML structures
        html_content = md
        # Replace headers
        html_content = re_replace_md(html_content)
        
        html_skeleton = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{title}</title>
            <style>
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
                    line-height: 1.6;
                    color: #1e293b;
                    max-width: 800px;
                    margin: 40px auto;
                    padding: 0 20px;
                    background-color: #f8fafc;
                }}
                h1 {{
                    color: #4f46e5;
                    border-bottom: 2px solid #e2e8f0;
                    padding-bottom: 10px;
                    font-size: 2.2rem;
                }}
                h2 {{
                    color: #1e1b4b;
                    margin-top: 30px;
                    border-bottom: 1px solid #e2e8f0;
                    padding-bottom: 6px;
                }}
                h3 {{
                    color: #6366f1;
                }}
                code {{
                    background-color: #e2e8f0;
                    padding: 2px 6px;
                    border-radius: 4px;
                    font-family: monospace;
                }}
                pre {{
                    background-color: #1e293b;
                    color: #f8fafc;
                    padding: 15px;
                    border-radius: 8px;
                    overflow-x: auto;
                }}
                blockquote {{
                    border-left: 4px solid #6366f1;
                    padding-left: 15px;
                    margin-left: 0;
                    color: #475569;
                    font-style: italic;
                }}
                details {{
                    background-color: #f1f5f9;
                    padding: 10px 15px;
                    border-radius: 6px;
                    margin: 10px 0;
                    border: 1px solid #cbd5e1;
                }}
                summary {{
                    font-weight: bold;
                    cursor: pointer;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        return html_skeleton

    @staticmethod
    def export_to_docx(title: str, content_dict: dict) -> bytes:
        """Generates a formatted Word Document and returns its bytes."""
        doc = DocxDoc()
        
        # Style Title
        t = doc.add_heading(level=0)
        t_run = t.add_run(title)
        t_run.font.name = 'Arial'
        t_run.font.color.rgb = RGBColor(79, 70, 229) # Indigo
        
        for section, content in content_dict.items():
            h = doc.add_heading(level=1)
            h_run = h.add_run(section)
            h_run.font.name = 'Arial'
            h_run.font.color.rgb = RGBColor(30, 27, 75)
            
            # Clean content lines
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                    
                # Basic Markdown conversion inside Word
                if para.startswith("###"):
                    sub_h = doc.add_heading(level=2)
                    sub_run = sub_h.add_run(para.replace("###", "").strip())
                    sub_run.font.name = 'Arial'
                    sub_run.font.color.rgb = RGBColor(99, 102, 241)
                elif para.startswith("- ") or para.startswith("* "):
                    lines = para.split("\n")
                    for line in lines:
                        item = line.replace("- ", "").replace("* ", "").strip()
                        doc.add_paragraph(item, style='List Bullet')
                else:
                    doc.add_paragraph(para)
                    
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    @staticmethod
    def export_to_pdf(title: str, content_dict: dict) -> bytes:
        """Generates a clean, readable PDF report and returns its bytes."""
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # UTF-8 core font mappings
        pdf.set_font("helvetica", "B", 18)
        pdf.set_text_color(79, 70, 229)
        pdf.cell(0, 15, title, ln=True, align="C")
        pdf.ln(10)
        
        for section, content in content_dict.items():
            # Section Header
            pdf.set_font("helvetica", "B", 14)
            pdf.set_text_color(30, 27, 75)
            # Use multi_cell to wrap long text
            pdf.multi_cell(0, 10, section)
            pdf.ln(2)
            
            # Section Body
            pdf.set_font("helvetica", "", 10)
            pdf.set_text_color(51, 65, 85)
            
            # Clean markdown artifacts from pdf display
            cleaned_content = content.replace("### ", "").replace("**", "")
            
            # Process paragraph by paragraph
            paragraphs = cleaned_content.split("\n\n")
            for p in paragraphs:
                p = p.strip()
                if not p:
                    continue
                # encode characters to match pdf latin-1 defaults
                safe_text = p.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 6, safe_text)
                pdf.ln(3)
                
            pdf.ln(5)
            
        # Get output as bytes
        pdf_bytes = pdf.output()
        return bytes(pdf_bytes)

def re_replace_md(md_str: str) -> str:
    """Helper to convert basic markdown text to clean HTML tags."""
    # Replace headings
    out = md_str
    out = re_sub_safe(r'^###\s+(.*)$', r'<h3>\1</h3>', out)
    out = re_sub_safe(r'^##\s+(.*)$', r'<h2>\1</h2>', out)
    out = re_sub_safe(r'^#\s+(.*)$', r'<h1>\1</h1>', out)
    
    # Replace bold/italics
    out = re_sub_safe(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', out)
    out = re_sub_safe(r'\*(.*?)\*', r'<em>\1</em>', out)
    
    # Replace paragraphs
    paras = out.split("\n\n")
    cleaned_paras = []
    for p in paras:
        p = p.strip()
        if not p:
            continue
        if p.startswith("<h") or p.startswith("<pre") or p.startswith("<details") or p.startswith("<blockquote"):
            cleaned_paras.append(p)
        else:
            # Replace line breaks inside paragraph
            p = p.replace("\n", "<br>")
            cleaned_paras.append(f"<p>{p}</p>")
            
    return "\n".join(cleaned_paras)

def re_sub_safe(pattern, repl, text):
    import re
    return re.sub(pattern, repl, text, flags=re.MULTILINE)
